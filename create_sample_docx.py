#!/usr/bin/env python3
"""
Create a sample Word document.
"""

from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).parent))

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Please install python-docx first: pip install python-docx")
    sys.exit(1)

# Create the document
doc = Document()

# Title
title = doc.add_heading('法律意见书示例', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Recipient
p = doc.add_paragraph()
p.add_run('致：').bold = True
p.add_run('北京示例科技有限公司')

doc.add_paragraph()

# Body
doc.add_heading('一、当事人信息', level=1)

doc.add_paragraph('1. 公司基本情况')
p = doc.add_paragraph()
p.add_run('   - 公司名称：').bold = True
p.add_run('北京示例科技有限公司')

p = doc.add_paragraph()
p.add_run('   - 统一社会信用代码：').bold = True
p.add_run('91110000000000000X')

p = doc.add_paragraph()
p.add_run('   - 法定代表人：').bold = True
p.add_run('张三')

p = doc.add_paragraph()
p.add_run('   - 地址：').bold = True
p.add_run('北京市海淀区中关村大街1号')

doc.add_paragraph()
doc.add_paragraph('2. 相关人员')
p = doc.add_paragraph()
p.add_run('   - 董事长：').bold = True
p.add_run('张三')

p = doc.add_paragraph()
p.add_run('   - 总经理：').bold = True
p.add_run('李四')

doc.add_heading('二、联系方式', level=1)

p = doc.add_paragraph()
p.add_run('- 公司电话：').bold = True
p.add_run('0755-88888888')

p = doc.add_paragraph()
p.add_run('- 移动电话：').bold = True
p.add_run('13812345678')

p = doc.add_paragraph()
p.add_run('- 电子邮箱：').bold = True
p.add_run('contact@example.com')

doc.add_heading('三、身份证信息（示例）', level=1)

p = doc.add_paragraph()
p.add_run('- 张三：').bold = True
p.add_run('110101199001011234')

p = doc.add_paragraph()
p.add_run('- 李四：').bold = True
p.add_run('110101199002025678')

doc.add_heading('四、日期信息', level=1)

p = doc.add_paragraph()
p.add_run('- 本意见书出具日期：').bold = True
p.add_run('2026年2月27日')

# Save the document
output_path = Path(__file__).parent / 'examples' / 'sample.docx'
doc.save(str(output_path))

print(f"Sample Word document created: {output_path}")
