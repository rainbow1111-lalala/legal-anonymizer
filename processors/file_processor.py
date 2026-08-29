"""
File Processor - Handles reading/writing various file formats
文件处理器 - 处理各种文件格式的读写
"""

import os
import sys
import json
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any

# 尝试导入可选依赖
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    from rapidocr import RapidOCR as _RapidOCR
    HAS_RAPIDOCR = True
except ImportError:
    HAS_RAPIDOCR = False


def _io_bytes_png(pil_image) -> bytes:
    """PIL Image → PNG bytes，避免每次写临时文件"""
    import io as _io
    buf = _io.BytesIO()
    pil_image.save(buf, format='PNG')
    return buf.getvalue()

try:
    from paddleocr import PaddleOCR as _PaddleOCR
    HAS_PADDLEOCR = True
except ImportError:
    HAS_PADDLEOCR = False

try:
    import pytesseract
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

# 至少有一种 OCR 引擎可用
HAS_OCR = HAS_PIL and (HAS_RAPIDOCR or HAS_PADDLEOCR or HAS_TESSERACT)


class FileProcessor:
    """文件处理器"""

    # 引擎单例缓存
    _rapidocr_instance = None
    _paddle_ocr_instance = None

    @classmethod
    def _get_rapidocr(cls):
        """懒加载 RapidOCR 实例（默认引擎，轻量、快速）"""
        if cls._rapidocr_instance is None and HAS_RAPIDOCR:
            cls._rapidocr_instance = _RapidOCR()
        return cls._rapidocr_instance

    @classmethod
    def _get_paddle_ocr(cls):
        """懒加载 PaddleOCR 实例（精准模式，慢但对复杂排版更强）"""
        if cls._paddle_ocr_instance is None and HAS_PADDLEOCR:
            # PaddleOCR 3.5 使用 mobile 模型（server 模型 CPU 推理太慢）
            cls._paddle_ocr_instance = _PaddleOCR(
                lang='ch',
                use_textline_orientation=True,
                text_detection_model_name='PP-OCRv5_mobile_det',
                text_recognition_model_name='PP-OCRv5_mobile_rec',
            )
        return cls._paddle_ocr_instance

    def __init__(self):
        self.supported_formats = {
            'input': ['.txt', '.md', '.pdf', '.docx', '.doc'],
            'output': ['.txt', '.md', '.pdf', '.docx']
        }

    def extract_text(self, file_path: str, use_ocr: bool = False, ocr_engine: str = 'rapidocr',
                     progress_callback=None) -> str:
        """
        从文件中提取文本

        Args:
            file_path: 文件路径
            use_ocr: 是否使用OCR（仅适用于PDF/图片）
            ocr_engine: 'rapidocr'（默认，快）| 'paddleocr'（慢但精准）| 'tesseract'（兜底）
            progress_callback: 可选的进度回调，签名 fn(stage:str, current:int, total:int)
                              stage 可为 'ocr_page_done' / 'extract_done'

        Returns:
            提取的文本内容
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = path.suffix.lower()

        if suffix == '.pdf':
            text = self._extract_pdf_text(file_path, use_ocr, ocr_engine, progress_callback=progress_callback)
        elif suffix == '.docx':
            text = self._extract_docx_text(file_path)
        elif suffix == '.doc':
            text = self._extract_doc_text(file_path)
        elif suffix in ['.txt', '.md']:
            text = self._extract_plain_text(file_path)
        elif suffix in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif', '.webp']:
            text = self._extract_image_text(file_path, ocr_engine)
        else:
            # 默认尝试作为文本文件读取
            try:
                text = self._extract_plain_text(file_path)
            except:
                raise ValueError(f"不支持的文件格式: {suffix}")

        # 中文字间换行合并：PDF 抽取时排版导致的中文字符间换行（如"XX\n公司"、"姓\n名"）
        # 会把同一实体拆成多个片段，规范化后再交给检测层
        return self._normalize_cjk_linebreaks(text)

    @staticmethod
    def _normalize_cjk_linebreaks(text: str) -> str:
        """
        1. 去掉 XML 不允许的控制字符（NULL 字节等）—— OCR 输出常见
        2. 合并两个中文字符之间的单个换行符（保留段落换行=连续换行）
        """
        import re
        # 先清除 XML 非法控制字符，防止下游 docx/PDF 写入失败
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        # 单个 \n（前后都是中文），替换为空
        text = re.sub(r'([一-龥])\n([一-龥])', r'\1\2', text)
        # 中文+\n+中文标点
        text = re.sub(r'([一-龥])\n([（）《》、，。；：])', r'\1\2', text)
        text = re.sub(r'([（《、，。；：])\n([一-龥])', r'\1\2', text)
        # 数字与中文互相粘接的换行（地址"宿舍6 栋\n102"、案号"（2006）192\n号"）
        text = re.sub(r'([一-龥])\n(\d)', r'\1\2', text)
        text = re.sub(r'(\d)\n([一-龥])', r'\1\2', text)
        # 中文和英文字母之间的换行（"SOHO现代城A座1203\n室"这类）
        text = re.sub(r'([一-龥])\n([A-Za-z])', r'\1\2', text)
        text = re.sub(r'([A-Za-z])\n([一-龥])', r'\1\2', text)
        return text

    def _extract_pdf_text(self, pdf_path: str, use_ocr: bool = False, ocr_engine: str = 'rapidocr',
                          progress_callback=None) -> str:
        """从PDF提取文本。ocr_engine: rapidocr(默认) | paddleocr | tesseract"""
        if not HAS_PYMUPDF:
            raise ImportError("需要安装 PyMuPDF: pip install pymupdf")

        import time as _time
        from concurrent.futures import ThreadPoolExecutor

        doc = fitz.open(pdf_path)
        page_count = doc.page_count

        # 先决定每一页走 OCR 还是直接抽取（PDF 渲染必须在主线程，PyMuPDF Page 不是线程安全的）
        page_jobs = []  # [(page_num, native_text, ocr_image_or_None)]
        for page_num, page in enumerate(doc, 1):
            native_text = page.get_text()
            need_ocr = use_ocr or (HAS_OCR and len(native_text.strip()) < 50)
            # use_ocr=True 但页面已有充足文字层时，不强制 OCR（省 5-10x 时间）
            if use_ocr and len(native_text.strip()) >= 200:
                need_ocr = False

            ocr_img = None
            if need_ocr and HAS_PIL:
                # 主线程渲染像素图（150 DPI）
                pix = page.get_pixmap(matrix=fitz.Matrix(2.08, 2.08))
                ocr_img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                pix = None  # 早释放
            page_jobs.append((page_num, native_text, ocr_img))

        doc.close()  # 释放 PDF 资源，OCR 阶段不再需要

        ocr_pages = sum(1 for _, _, img in page_jobs if img is not None)
        if ocr_pages > 0:
            print(f"[OCR] 共 {page_count} 页，需要 OCR 的 {ocr_pages} 页（其余直接抽取文字层）", flush=True)

        # 多线程并行：ONNX Runtime 内部已多线程，2 worker 已经能让 CPU 跑满
        # 太多 worker 反而内存压力大且效益递减
        max_workers = min(2, ocr_pages or 1)

        # 用计数器，多线程下也安全（GIL 保护 int += 1）
        ocr_done_counter = [0]
        ocr_total = ocr_pages

        def process_page(job):
            page_num, native_text, ocr_img = job
            t0 = _time.time()
            if ocr_img is not None:
                text = self._ocr_image(ocr_img, ocr_engine)
                text = self._fix_ocr_text(text)
                elapsed = _time.time() - t0
                print(f"[OCR] 第 {page_num}/{page_count} 页完成，耗时 {elapsed:.1f}s", flush=True)
                ocr_done_counter[0] += 1
                if progress_callback:
                    try:
                        progress_callback('ocr_page_done', ocr_done_counter[0], ocr_total)
                    except Exception:
                        pass
            else:
                text = self._fix_line_breaks(native_text)
            return page_num, text

        text_parts = [None] * page_count
        if max_workers > 1 and ocr_pages > 1:
            with ThreadPoolExecutor(max_workers=max_workers) as pool:
                for page_num, text in pool.map(process_page, page_jobs):
                    text_parts[page_num - 1] = text
        else:
            for job in page_jobs:
                page_num, text = process_page(job)
                text_parts[page_num - 1] = text

        return '\n\n'.join(text_parts)

    def _ocr_pdf_page(self, page, ocr_engine: str = 'rapidocr') -> str:
        """对 PDF 页面做 OCR。先渲染成图像，再调引擎"""
        if not HAS_PIL:
            return ""

        try:
            # 150 DPI 渲染（法律文书清晰度足够，比 200 DPI 快约 30%）
            pix = page.get_pixmap(matrix=fitz.Matrix(2.08, 2.08))  # 2.08 × 72 ≈ 150
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            return self._ocr_image(img, ocr_engine)
        except Exception:
            return ""

    def _ocr_image(self, img, ocr_engine: str = 'rapidocr') -> str:
        """对 PIL Image 做 OCR，按指定引擎优先级尝试"""
        # 按指定引擎尝试
        engines_order = []
        if ocr_engine == 'paddleocr':
            engines_order = ['paddleocr', 'rapidocr', 'tesseract']
        elif ocr_engine == 'tesseract':
            engines_order = ['tesseract', 'rapidocr', 'paddleocr']
        else:  # 默认 rapidocr
            engines_order = ['rapidocr', 'paddleocr', 'tesseract']

        for eng in engines_order:
            try:
                if eng == 'rapidocr' and HAS_RAPIDOCR:
                    return self._ocr_with_rapidocr(img)
                if eng == 'paddleocr' and HAS_PADDLEOCR:
                    return self._ocr_with_paddleocr(img)
                if eng == 'tesseract' and HAS_TESSERACT:
                    return self._ocr_with_tesseract(img)
            except Exception:
                continue

        return ""

    def _ocr_with_rapidocr(self, img) -> str:
        """RapidOCR 实现：快、轻量、默认选择
        优化：直接传 numpy 数组，跳过 PNG 编码（每页省 200-500ms）
        """
        import numpy as np
        ocr = self._get_rapidocr()
        # PIL Image -> numpy（RapidOCR 直接接受 numpy 数组）
        img_array = np.array(img)
        result = ocr(img_array)
        if result and result.txts:
            lines = []
            for txt, score in zip(result.txts, result.scores or []):
                if score > 0.5:
                    lines.append(txt)
            return '\n'.join(lines) if lines else '\n'.join(result.txts)
        return ""

    def _ocr_with_paddleocr(self, img) -> str:
        """PaddleOCR 3.5 实现：慢但对复杂排版更准"""
        import numpy as np
        ocr = self._get_paddle_ocr()
        img_array = np.array(img)
        result = ocr.predict(img_array)
        if result:
            lines = []
            for page_result in result:
                # PaddleOCR 3.5 返回 dict-like 结果，含 rec_texts / rec_scores
                rec_texts = page_result.get('rec_texts') if hasattr(page_result, 'get') else None
                rec_scores = page_result.get('rec_scores') if hasattr(page_result, 'get') else None
                if rec_texts:
                    if rec_scores:
                        for t, s in zip(rec_texts, rec_scores):
                            if s > 0.6:
                                lines.append(t)
                    else:
                        lines.extend(rec_texts)
            return '\n'.join(lines)
        return ""

    def _ocr_with_tesseract(self, img) -> str:
        """Tesseract 兜底"""
        custom_config = r'-l chi_sim+eng --oem 1 --psm 6'
        return pytesseract.image_to_string(img, config=custom_config)

    def _fix_ocr_text(self, text: str) -> str:
        """修复OCR提取文本的断行和空格问题，使句子连贯"""
        import re

        # 第一步：去除OCR噪音字符（扫描边缘产生的竖线、乱码等）
        # 去掉行尾/行首的 | 和其他常见OCR噪音
        text = re.sub(r'\s*\|\s*', '', text)
        # 去掉单独的乱码字符（非中文、非ASCII字母数字、非标点的孤立字符）
        # 保留中文、英文字母、数字、常见标点
        text = re.sub(r'(?<=[。，、；：！？）\)）])\s*[a-zA-Z]{1,3}\s*(?=\n|$)', '', text)

        # 第二步：清理多余空白
        # 压缩连续空格为单个
        text = re.sub(r'[ \t]{2,}', ' ', text)

        # 第三步：去掉中文字符之间的空格
        # 多轮处理确保全部清除
        for _ in range(3):
            text = re.sub(r'([\u4e00-\u9fa5])\s+([\u4e00-\u9fa5])', r'\1\2', text)
        # 中文和数字之间的多余空格
        text = re.sub(r'([\u4e00-\u9fa5])\s+(\d)', r'\1\2', text)
        text = re.sub(r'(\d)\s+([\u4e00-\u9fa5])', r'\1\2', text)
        # 中文和标点之间的空格
        text = re.sub(r'([\u4e00-\u9fa5])\s+([，。、；：！？）\)》」』】])', r'\1\2', text)
        text = re.sub(r'([（\(《「『【])\s+([\u4e00-\u9fa5])', r'\1\2', text)

        # 第四步：合并断行
        lines = text.split('\n')
        merged_lines = []
        buffer = ''

        for line in lines:
            stripped = line.strip()

            # 跳过空行
            if not stripped:
                if buffer:
                    merged_lines.append(buffer)
                    buffer = ''
                continue

            # 跳过纯噪音行（全是非中文非数字的短行）
            clean_chars = re.sub(r'[^a-zA-Z0-9\u4e00-\u9fa5]', '', stripped)
            if len(clean_chars) < 2:
                continue

            if not buffer:
                buffer = stripped
            else:
                last_char = buffer[-1] if buffer else ''

                # 只在明确的句子结束标点处分段
                if last_char in '。！？':
                    merged_lines.append(buffer)
                    buffer = stripped
                else:
                    # 其他情况全部合并（包括逗号、分号结尾）
                    buffer += stripped

        if buffer:
            merged_lines.append(buffer)

        return '\n'.join(merged_lines)

    def _fix_line_breaks(self, text: str) -> str:
        """修复PDF文本提取中常见的换行分割问题（非OCR模式）"""
        import re

        # 修复组织机构名称被换行分割（通用规则）
        org_suffixes = (
            r'有限公司|有限责任公司|股份有限公司|股份公司'
            r'|集团公司|集团|律师事务所|会计师事务所|公证处'
            r'|人民法院|人民检察院|人民政府'
            r'|公安局|派出所|管理局|监督局|委员会'
            r'|居委会|村委会|街道办事处|办事处'
            r'|研究院|研究所|实验室|大学|学院|医院'
            r'|银行|信用社|基金会|协会|商会|学会'
        )
        text = re.sub(
            rf'([\u4e00-\u9fa5]{{1,15}})\n({org_suffixes})',
            r'\1\2', text
        )

        # 修复公司名在后缀之前断行的情况
        # 例如 "国新健康保障服\n务集团股份有限公司" → 拼接
        # 不加严格的后向断言，因为 PDF 签章页等格式后面可能直接跟"年月日"
        text = re.sub(
            rf'([\u4e00-\u9fa5])'
            rf'\n'
            rf'([\u4e00-\u9fa5]{{1,20}}(?:{org_suffixes}))',
            r'\1\2', text
        )

        # 修复常见的词语被换行分割
        text = re.sub(
            r'([\u4e00-\u9fa5])(?<![。！？；：，、）》」』】])\n(?=[\u4e00-\u9fa5]{1,3}(?:[。！？；：，、）》」』】\n]|$))',
            r'\1', text
        )

        return text

    def _extract_doc_text(self, doc_path: str) -> str:
        """从旧版 .doc 文件提取文本"""
        import subprocess
        import platform

        # macOS: 使用 textutil
        if platform.system() == 'Darwin':
            try:
                result = subprocess.run(
                    ['textutil', '-convert', 'txt', '-stdout', doc_path],
                    capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout
            except (subprocess.TimeoutExpired, FileNotFoundError):
                pass

        # Linux: 尝试 antiword
        try:
            result = subprocess.run(
                ['antiword', doc_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass

        # Linux: 尝试 libreoffice
        try:
            import tempfile
            with tempfile.TemporaryDirectory() as tmpdir:
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'txt:Text', '--outdir', tmpdir, doc_path],
                    capture_output=True, text=True, timeout=60
                )
                if result.returncode == 0:
                    txt_file = Path(tmpdir) / (Path(doc_path).stem + '.txt')
                    if txt_file.exists():
                        return txt_file.read_text(encoding='utf-8')
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass

        raise ImportError(
            "无法读取 .doc 文件。请安装以下工具之一：\n"
            "  macOS: textutil（系统自带）\n"
            "  Linux: sudo apt install antiword  或  sudo apt install libreoffice\n"
            "  或者将文件转换为 .docx 格式后重试"
        )

    def _extract_docx_text(self, docx_path: str) -> str:
        """从Word文档提取文本（过滤HYPERLINK等域代码）"""
        if not HAS_DOCX:
            raise ImportError("需要安装 python-docx: pip install python-docx")

        import re

        doc = Document(docx_path)
        content = []

        def clean_hyperlinks(t):
            t = re.sub(r'HYPERLINK\s+"[^"]*"\s*(?:\\[a-z]\s+"[^"]*"\s*)*', '', t)
            t = re.sub(r'HYPERLINK\s+\S+\s*', '', t)
            return t

        def extract_paragraphs(paras):
            for para in paras:
                text = clean_hyperlinks(para.text)
                if text.strip():
                    content.append(text)

        # 页眉页脚（各 section）
        for section in doc.sections:
            for hf in [section.header, section.footer,
                       section.first_page_header, section.first_page_footer,
                       section.even_page_header, section.even_page_footer]:
                if hf and not hf.is_linked_to_previous:
                    extract_paragraphs(hf.paragraphs)

        # 正文段落
        extract_paragraphs(doc.paragraphs)

        # 表格
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    text = clean_hyperlinks(cell.text.strip())
                    if text:
                        row_text.append(text)
                if row_text:
                    content.append(' | '.join(row_text))

        return '\n'.join(content)

    def anonymize_docx_inplace(self, input_path: str, output_path: str, mapping: dict) -> bool:
        """
        在保留原始Word文档所有格式的基础上进行脱敏替换。

        通过直接操作 XML 层的 <w:t> 元素实现，覆盖：
        - 正文段落、表格、文本框
        - 超链接内文本
        - 页眉页脚
        - 批注 (comments)
        - 修订/追踪更改 (track changes: <w:ins>, <w:del>)
        - 脚注、尾注

        所有 run 的字体、字号、颜色、加粗、段落间距等格式属性完全保留。
        """
        if not HAS_DOCX:
            return False

        try:
            from lxml import etree
        except ImportError:
            # python-docx 自带 lxml 依赖
            return False

        try:
            doc = Document(input_path)

            # 构建替换表（按长度降序，避免子串问题）
            replacements = {}
            for (etype, original), masked in mapping.items():
                replacements[original] = masked
            sorted_originals = sorted(replacements.keys(), key=len, reverse=True)

            if not sorted_originals:
                doc.save(output_path)
                return True

            W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            W_T = f'{{{W_NS}}}t'
            W_R = f'{{{W_NS}}}r'
            W_P = f'{{{W_NS}}}p'
            XML_NS = 'http://www.w3.org/XML/1998/namespace'

            def get_t_elements_from_para(para_elem):
                """获取段落中所有 <w:t> 元素（包括 w:ins/w:del/w:hyperlink 内部的）"""
                return list(para_elem.iter(W_T))

            def replace_in_paragraph(para_elem):
                """对单个段落做所有替换，逐个替换并刷新状态"""
                for original in sorted_originals:
                    masked = replacements[original]
                    # 每次替换后重新获取 t 元素（因为文本已变化）
                    while True:
                        t_elems = get_t_elements_from_para(para_elem)
                        if not t_elems:
                            break
                        texts = [t.text or '' for t in t_elems]
                        full_text = ''.join(texts)

                        pos = full_text.find(original)
                        if pos == -1:
                            break

                        end_pos = pos + len(original)

                        # 定位到具体的 <w:t> 元素和偏移
                        char_idx = 0
                        start_ti = end_ti = -1
                        start_offset = end_offset = 0

                        for ti, text in enumerate(texts):
                            t_start = char_idx
                            t_end = char_idx + len(text)

                            if start_ti == -1 and t_start <= pos < t_end:
                                start_ti = ti
                                start_offset = pos - t_start

                            if t_start < end_pos <= t_end:
                                end_ti = ti
                                end_offset = end_pos - t_start
                                break

                            char_idx = t_end

                        if start_ti == -1 or end_ti == -1:
                            break

                        if start_ti == end_ti:
                            # 同一个 <w:t> 内替换
                            t = t_elems[start_ti]
                            t.text = t.text[:start_offset] + masked + t.text[end_offset:]
                        else:
                            # 跨 <w:t> 替换：替换文本放入第一个，中间清空，尾部保留
                            t_elems[start_ti].text = texts[start_ti][:start_offset] + masked
                            for ti in range(start_ti + 1, end_ti):
                                t_elems[ti].text = ''
                            t_elems[end_ti].text = texts[end_ti][end_offset:]

                        # 确保含前后空格的文本保留 xml:space="preserve"
                        for ti in range(start_ti, min(end_ti + 1, len(t_elems))):
                            t = t_elems[ti]
                            if t.text and (t.text[0] == ' ' or t.text[-1] == ' ' or '  ' in t.text):
                                t.set(f'{{{XML_NS}}}space', 'preserve')

            def process_element_tree(root_elem):
                """处理一个 XML 元素树中的所有段落"""
                for para in root_elem.iter(W_P):
                    replace_in_paragraph(para)

            # 1. 处理主文档（正文、表格、文本框、超链接、修订记录等全部覆盖）
            process_element_tree(doc.element.body)

            # 2. 处理页眉页脚
            for section in doc.sections:
                for hf in [section.header, section.footer,
                           section.first_page_header, section.first_page_footer,
                           section.even_page_header, section.even_page_footer]:
                    try:
                        if hf and hf.is_linked_to_previous is False:
                            process_element_tree(hf.element)
                    except Exception:
                        pass

            # 3. 处理批注、脚注、尾注等关联部件
            PART_REL_TYPES = [
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments',
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes',
                'http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes',
            ]
            for rel in doc.part.rels.values():
                if rel.reltype in PART_REL_TYPES:
                    try:
                        part_elem = rel.target_part.element
                        process_element_tree(part_elem)
                    except Exception:
                        pass

            doc.save(output_path)
            return True

        except Exception:
            import traceback
            traceback.print_exc()
            return False

    def restore_docx_inplace(self, input_path: str, output_path: str, mapping_data: dict) -> bool:
        """在保留 Word 原始排版的前提下，将占位符反向替换为原文。

        ``mapping_data`` 可以是完整字典，也可以是
        ``{placeholder: {type, original}}`` 或 ``{placeholder: original}``。
        """
        if not isinstance(mapping_data, dict):
            return False
        if isinstance(mapping_data.get('mapping'), dict):
            mapping_data = mapping_data['mapping']

        reverse_mapping = {}
        for placeholder, info in mapping_data.items():
            if isinstance(info, dict):
                original = info.get('original', '')
            else:
                original = str(info or '')
            if placeholder and original:
                reverse_mapping[('restore', placeholder)] = original

        return self.anonymize_docx_inplace(input_path, output_path, reverse_mapping)

    def _extract_plain_text(self, file_path: str) -> str:
        """从纯文本文件提取（自动尝试常见中文编码）"""
        for enc in ('utf-8-sig', 'utf-8', 'gbk', 'gb18030'):
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        # 最终兜底：替换无法解码的字符
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()

    def _extract_image_text(self, image_path: str, ocr_engine: str = 'rapidocr') -> str:
        """从图片提取文本（OCR）"""
        if not HAS_OCR:
            raise ImportError(
                "需要安装 OCR 依赖:\n"
                "  推荐: pip install rapidocr pillow\n"
                "  备选: pip install paddleocr paddlepaddle\n"
                "  兜底: pip install pytesseract（需另装 tesseract 二进制）"
            )

        img = Image.open(image_path)
        return self._ocr_image(img, ocr_engine)

    def anonymize_pdf_inplace(
        self, input_path: str, output_path: str, mapping: dict,
        whitebox_only: bool = False,
    ) -> bool:
        """
        PDF → PDF 原地脱敏（保留原 PDF 的全部格式：字体、排版、盖章、签名）

        用 PyMuPDF 的 redact 机制：
          1. 用 page.search_for 定位每个原始敏感文本的坐标框
          2. add_redact_annot 在原位置覆盖为白底 + 占位符文字
          3. apply_redactions 真正擦除

        中文字体处理：
          - 优先用 "china-s"（PyMuPDF 内置简体）
          - 失败回退系统 PingFang/STHeiti
          - 最后兜底 helv（英文字体，会丢失中文占位符的字形）

        Args:
            input_path: 源 PDF
            output_path: 脱敏后 PDF
            mapping: {(type, original): placeholder, ...}（与 TextMasker.mapping 结构相同）
        """
        if not HAS_PYMUPDF:
            return False

        try:
            import fitz as _fitz

            # 构建 original → masked 的替换字典，按长度降序避免子串问题
            replacements = {}
            for (etype, original), masked in mapping.items():
                if original and masked:
                    replacements[original] = masked
            if not replacements:
                # 无实体则直接复制
                doc = _fitz.open(input_path)
                doc.save(output_path)
                doc.close()
                return True

            sorted_originals = sorted(replacements.keys(), key=len, reverse=True)

            # PyMuPDF 内置 china-s（简体）CJK 字体，无需 fontfile
            fontname = 'china-s'

            doc = _fitz.open(input_path)

            for page in doc:
                # 构建字符级 bbox 映射：允许 entity 跨行匹配
                flat_text, char_rects = self._build_char_map(page, _fitz)

                # 取页面正文的中位字号（让占位符字号跟原文一致，不再小一截）
                page_font_sizes = []
                try:
                    pd = page.get_text("dict")
                    for block in pd.get("blocks", []):
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                size = span.get("size", 0)
                                if size and len(span.get("text", "")) > 0:
                                    page_font_sizes.append(size)
                except Exception:
                    pass
                page_font_sizes.sort()
                page_default_size = (
                    page_font_sizes[len(page_font_sizes) // 2]
                    if page_font_sizes else 11
                )

                first_seen: set = set()  # 同一个 original 只替换一次占位符，余下清空

                for original in sorted_originals:
                    masked = replacements[original]
                    occurrences = self._find_rects_for_entity(
                        flat_text, char_rects, original, _fitz
                    )
                    for occ_idx, rects in enumerate(occurrences):
                        if not rects:
                            continue
                        # 多个 rect（跨行情况）：占位符只写在第一段；其余擦除
                        for ri, rect in enumerate(rects):
                            # 轻微扩大 rect 以保证完全覆盖字形（字符 bbox 有时不含 descender）
                            padded = _fitz.Rect(
                                rect.x0 - 0.5, rect.y0 - 1.5,
                                rect.x1 + 0.5, rect.y1 + 1.5,
                            )
                            # 优先用页面正文的中位字号，兜底用 rect.height 估算
                            fontsize = page_default_size
                            # 如果原 rect 高度比页字号小很多（小字段），用 rect 估算更合适
                            est = rect.height * 0.85
                            if est < page_default_size * 0.7:
                                fontsize = max(8, est)
                            page.add_redact_annot(
                                padded,
                                text="" if whitebox_only else (masked if ri == 0 else ""),
                                fontname=fontname,
                                fontsize=fontsize,
                                text_color=(0, 0, 0),
                                fill=(1, 1, 1),
                                align=_fitz.TEXT_ALIGN_LEFT,
                            )

                # 应用 redaction。
                #   images=PDF_REDACT_IMAGE_NONE → 不擦图片（保留盖章/签名）
                #   graphics=PDF_REDACT_LINE_ART_NONE → 不擦线条（保留表格框线）
                #   text=True → 擦除命中区的文字
                page.apply_redactions(
                    images=_fitz.PDF_REDACT_IMAGE_NONE,
                    graphics=_fitz.PDF_REDACT_LINE_ART_NONE,
                )

                # 收尾清扫：某些 PDF（如金格签章文件）有第二层文本，第一轮没擦到。
                # 用 search_for 找残留，最多清扫 2 轮。
                for _ in range(2):
                    had_residue = False
                    for original in sorted_originals:
                        masked = replacements[original]
                        rects = page.search_for(original)
                        for r in rects:
                            if r.width > 0.5:  # 跳过零宽度幽灵
                                had_residue = True
                                padded = _fitz.Rect(
                                    r.x0 - 0.5, r.y0 - 1.5,
                                    r.x1 + 0.5, r.y1 + 1.5,
                                )
                                fontsize = max(7, min(13, r.height * 0.72))
                                page.add_redact_annot(
                                    padded,
                                    text="" if whitebox_only else masked,
                                    fontname=fontname,
                                    fontsize=fontsize,
                                    text_color=(0, 0, 0),
                                    fill=(1, 1, 1),
                                    align=_fitz.TEXT_ALIGN_LEFT,
                                )
                    if not had_residue:
                        break
                    page.apply_redactions(
                        images=_fitz.PDF_REDACT_IMAGE_NONE,
                        graphics=_fitz.PDF_REDACT_LINE_ART_NONE,
                    )

            doc.save(output_path, deflate=True, garbage=3)
            doc.close()
            return True
        except Exception:
            import traceback
            traceback.print_exc()
            return False

    @staticmethod
    def _build_char_map(page, _fitz):
        """
        用 rawdict 取页面中每个字符的 bbox。
        返回 (flat_text, char_rects)，两者一一对应；行尾追加 '\\n' 且 rect=None。

        关键：char bbox 的 y 坐标只覆盖 glyph 原点附近，会漏擦 ascender/descender。
        这里把每个字符的 y 坐标扩展为所在 line 的完整 bbox y 范围，确保 apply_redactions
        能把整个字形擦除干净。
        """
        raw = page.get_text("rawdict")
        flat = []
        rects = []
        for block in raw.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                line_bbox = line.get("bbox")
                if line_bbox:
                    line_y0, line_y1 = line_bbox[1], line_bbox[3]
                else:
                    line_y0 = line_y1 = None
                for span in line.get("spans", []):
                    for ch in span.get("chars", []):
                        c = ch.get("c") or ""
                        bbox = ch.get("bbox")
                        if not c:
                            continue
                        if bbox and line_y0 is not None:
                            # x 用 char 自己的，y 用 line 的（确保覆盖完整字形高度）
                            r = _fitz.Rect(bbox[0], line_y0, bbox[2], line_y1)
                        elif bbox:
                            r = _fitz.Rect(bbox)
                        else:
                            r = None
                        flat.append(c)
                        rects.append(r)
                flat.append("\n")
                rects.append(None)
        return "".join(flat), rects

    @staticmethod
    def _find_rects_for_entity(flat_text: str, char_rects: list, entity: str, _fitz):
        """
        在 flat_text 中找 entity 的所有出现位置，允许中文/数字之间有 \\n（排版换行）。
        返回 list of list：每个出现对应的 rect 列表（跨行则含多段）。
        """
        import re
        # 为 entity 构造"容忍换行"的正则：在 CJK/数字字符之间允许插入一个 \n
        def is_break_allowed(ch):
            return ('一' <= ch <= '鿿') or ch.isdigit() or ch.isalpha() or ch in ' \t'

        pattern_parts = []
        for i, ch in enumerate(entity):
            pattern_parts.append(re.escape(ch))
            if i < len(entity) - 1 and is_break_allowed(ch) and is_break_allowed(entity[i + 1]):
                pattern_parts.append(r"\n?")
        pattern = "".join(pattern_parts)

        occurrences = []
        for m in re.finditer(pattern, flat_text):
            rs = []
            current = None
            for i in range(m.start(), m.end()):
                r = char_rects[i] if i < len(char_rects) else None
                if r is None:
                    if current is not None:
                        rs.append(current)
                        current = None
                    continue
                if current is None:
                    current = _fitz.Rect(r)
                elif abs(r.y0 - current.y0) > 2 and abs(r.y1 - current.y1) > 2:
                    rs.append(current)
                    current = _fitz.Rect(r)
                else:
                    current |= r  # union
            if current is not None:
                rs.append(current)
            if rs:
                occurrences.append(rs)
        return occurrences

    def anonymize_scanned_pdf_inplace(
        self, input_path: str, output_path: str, mapping: dict,
        ocr_engine: str = 'rapidocr', whitebox_only: bool = False,
    ) -> bool:
        """
        扫描版 PDF 视觉脱敏（保留原页面图像，只把敏感字位置盖白底+占位符）

        与 anonymize_pdf_inplace 的区别：
          - 那个针对**文字层 PDF**：直接 redact 文字对象，保留所有原格式
          - 这个针对**扫描版 PDF**：原"文字"是图像内容，需要先 OCR 拿到字符位置，
            再用 PDF_REDACT_IMAGE_PIXELS 擦除底层图像并覆盖占位符

        效果：输出 PDF 看上去和原扫描版几乎一样（保留盖章/签名/页眉/纸张底色），
        只有敏感字处变成白色色块上的 "[PERSON_1]" 等占位符。

        Args:
            input_path: 源扫描 PDF
            output_path: 输出脱敏后 PDF
            mapping: TextMasker.mapping，{(etype, original): placeholder, ...}
            ocr_engine: OCR 引擎，目前仅支持 'rapidocr'（PaddleOCR 不返回 box 数据）
        """
        if not HAS_PYMUPDF or not HAS_PIL:
            return False
        if not HAS_RAPIDOCR:
            print("  ⚠️ 扫描版 PDF 视觉脱敏需要 RapidOCR，请安装：pip install rapidocr")
            return False

        try:
            import fitz as _fitz
            import numpy as np
            import time as _time

            # 保留 entity_type 信息（用于"独特识别部分"判断）
            replacements = {}
            entity_types = {}  # original -> type
            for (etype, original), masked in mapping.items():
                if original and masked and len(original) >= 1:
                    replacements[original] = masked
                    entity_types[original] = etype
            if not replacements:
                # 无敏感实体则直接复制
                doc = _fitz.open(input_path)
                doc.save(output_path)
                doc.close()
                return True

            # 长串优先（避免 "张三" 先匹配到 "张三李四" 的子串）
            sorted_originals = sorted(replacements.keys(), key=len, reverse=True)

            doc = _fitz.open(input_path)
            page_count = doc.page_count
            ocr = self._get_rapidocr()
            total_redacted = 0

            print(f"[视觉脱敏] 共 {page_count} 页扫描版 PDF，正在逐页处理...", flush=True)

            # 关键：用 PIL 直接在图像上画 redact，避免 PDF 旋转坐标转换问题
            from PIL import ImageDraw, ImageFont

            # 找系统中文字体（用于在 PIL 上绘占位符）
            font_path = None
            for cand in (
                '/System/Library/Fonts/PingFang.ttc',
                '/System/Library/Fonts/STHeiti Light.ttc',
                '/System/Library/Fonts/Hiragino Sans GB.ttc',
                'C:/Windows/Fonts/simhei.ttf',
                'C:/Windows/Fonts/msyh.ttc',
                '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
            ):
                if os.path.exists(cand):
                    font_path = cand
                    break

            # 输出新 PDF（每页用脱敏后的图像组装）
            out_doc = _fitz.open()

            for page_num, page in enumerate(doc, 1):
                t0 = _time.time()

                # 渲染 200 DPI（用作输出的图像，质量优先）
                render_dpi = 200
                scale = render_dpi / 72
                mat = _fitz.Matrix(scale, scale)
                pix = page.get_pixmap(matrix=mat)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img_array = np.array(img)
                pix = None

                # OCR
                result = ocr(img_array)
                if not result or not result.txts or result.boxes is None:
                    # 没识别到东西，原图直接放回
                    new_page = out_doc.new_page(width=page.rect.width, height=page.rect.height)
                    img_bytes = _io_bytes_png(img)
                    new_page.insert_image(new_page.rect, stream=img_bytes)
                    print(f"  第 {page_num}/{page_count} 页 OCR 无结果，原图放回", flush=True)
                    continue

                # 在 PIL 画布上作画
                draw = ImageDraw.Draw(img)

                # 关键：按"视觉行"分组排序后再拼接，避免 OCR 把同一行拆成左右两段后顺序错乱
                raw_lines = []
                for li, (line_text, line_box) in enumerate(zip(result.txts, result.boxes)):
                    bxs = [float(p[0]) for p in line_box]
                    bys = [float(p[1]) for p in line_box]
                    raw_lines.append({
                        'text': line_text,
                        'x0': min(bxs), 'x1': max(bxs),
                        'y0': min(bys), 'y1': max(bys),
                        'y_center': (min(bys) + max(bys)) / 2,
                        'h': max(bys) - min(bys),
                    })

                # 按 y_center 排，y 容差 = 平均行高 × 0.5 内视为同一视觉行
                raw_lines.sort(key=lambda L: L['y_center'])
                avg_h = sum(L['h'] for L in raw_lines) / max(len(raw_lines), 1)
                tolerance = max(avg_h * 0.5, 5)
                groups = []
                cur_group = []
                last_yc = None
                for L in raw_lines:
                    if last_yc is None or abs(L['y_center'] - last_yc) <= tolerance:
                        cur_group.append(L)
                    else:
                        groups.append(cur_group)
                        cur_group = [L]
                    last_yc = L['y_center']
                if cur_group:
                    groups.append(cur_group)
                # 每组内按 x0 排序（左→右），组间已按 y 排好
                ordered_lines = []
                for grp in groups:
                    grp.sort(key=lambda L: L['x0'])
                    ordered_lines.extend(grp)

                # 跨行匹配用拼接文本
                line_specs = []  # [(line_text, x0, x1, y0, y1, char_start_global)]
                joined = ""
                for L in ordered_lines:
                    char_start = len(joined)
                    line_specs.append((
                        L['text'], L['x0'], L['x1'], L['y0'], L['y1'], char_start,
                    ))
                    joined += L['text']

                # 全/半角等价规范化（OCR 容易把 ( 识为 (，反之亦然）
                def normalize(s):
                    return s.translate(str.maketrans('()[]【】《》＜＞，。；：！？',
                                                       '()[]<><><>,.;:!?'))
                joined_norm = normalize(joined)

                # === 策略：每行单独匹配实体的子串（避免依赖跨行拼接顺序）===
                # 对每个实体：遍历 OCR 每一行，找该行中包含的实体的最长子串。
                # 这样无论 OCR 怎么切分（横切、竖切、左右分段），任何在该行可见的
                # 实体片段都会被遮住，且 OCR 错位 / 顺序乱不影响结果。
                page_redacted = 0
                # 每行已遮区间 [(line_idx, char_start, char_end), ...]
                line_handled = {i: [] for i in range(len(line_specs))}
                # 哪些实体已经写过占位符（避免在多行重复写 [PERSON_1]）
                placeholder_written = set()

                def longest_common_substring(a, b, min_len=2):
                    """在 b 中找 a 的最长子串。返回 (a_start, b_start, length) 或 None"""
                    for L in range(min(len(a), len(b)), min_len - 1, -1):
                        for i in range(len(a) - L + 1):
                            sub = a[i:i + L]
                            j = b.find(sub)
                            if j != -1:
                                return (i, j, L)
                    return None

                def char_width_weight(c):
                    """估算字符宽度权重：CJK = 2，半角 = 1"""
                    if '一' <= c <= '鿿':
                        return 2.0  # CJK 汉字
                    if '　' <= c <= '〿' or '＀' <= c <= '￯':
                        return 2.0  # CJK 标点 / 全角
                    return 1.0  # 半角 ASCII / 数字 / 标点

                def pixel_offset_in_line(text, char_idx, line_x0, line_w):
                    """按 CJK 双倍宽权重，从 char_idx 字符位置算到像素 X"""
                    if not text:
                        return line_x0
                    weights = [char_width_weight(c) for c in text]
                    total = sum(weights)
                    if total <= 0:
                        return line_x0
                    cum = sum(weights[:char_idx])
                    return line_x0 + (cum / total) * line_w

                def get_distinctive_part(entity, etype):
                    """提取实体里的'品牌/独特识别部分'，去掉通用前后缀"""
                    import re as _re
                    if etype in ('company', 'law_firm', 'institution', 'bank_name'):
                        # 移除括号内容（含中英文括号）
                        s = _re.sub(r'[（(].*?[)）]', '', entity)
                        # 移除常见地名前缀
                        for p in ('北京', '上海', '广东省', '广州市', '深圳市', '深圳', '广州',
                                  '中国', '中华人民共和国', '中华', '广东', '浙江省', '浙江',
                                  '江苏省', '山东省', '河北省', '河南省'):
                            if s.startswith(p):
                                s = s[len(p):]
                        # 移除常见后缀（长后缀优先）
                        for sfx in sorted([
                            '律师事务所', '会计师事务所', '事务所',
                            '股份有限公司', '有限责任公司', '有限公司',
                            '集团有限公司', '集团公司',
                            '公司', '集团', '股份',
                        ], key=len, reverse=True):
                            if s.endswith(sfx):
                                s = s[:-len(sfx)]
                                break
                        return s.strip() if len(s.strip()) >= 2 else None
                    if etype == 'court':
                        for sfx in ('中级人民法院', '高级人民法院', '人民法院', '人民检察院',
                                    '法院', '检察院', '仲裁院', '仲裁委员会'):
                            if entity.endswith(sfx):
                                core = entity[:-len(sfx)]
                                return core if len(core) >= 2 else None
                        return entity
                    if etype == 'government':
                        for sfx in ('司法厅', '司法部', '人民政府', '管理委员会', '办公厅',
                                    '公安局', '派出所', '工商局'):
                            if entity.endswith(sfx):
                                core = entity[:-len(sfx)]
                                return core if len(core) >= 2 else None
                        return entity
                    # 人名/案号/信用代码 等：本身就是独特的
                    return entity

                for original in sorted_originals:
                    masked = replacements[original]
                    etype = entity_types.get(original, 'unknown')
                    orig_norm = normalize(original)
                    if len(orig_norm) < 2:
                        continue

                    # 数字/英文实体不要短匹配（容易误报）；中文最少 3 字以避免"公司"/"国际"等通用词误中
                    is_alnum = orig_norm.replace('.', '').replace('-', '').isalnum() and \
                               all(ord(c) < 128 for c in orig_norm)
                    if is_alnum:
                        min_match_len = 4
                    elif len(orig_norm) <= 3:
                        # 实体本身就只 2-3 字（人名/品牌）：必须完整匹配
                        min_match_len = len(orig_norm)
                    else:
                        # 长实体（4+ 字）：最少 3 字（避免"公司"等通用词误中，但 OCR 碎片化时仍能命中）
                        min_match_len = 3

                    # 计算"独特识别部分"（如 "北京XX（深圳）律师事务所" → "XX"）
                    distinctive = get_distinctive_part(original, etype)
                    distinctive_norm = normalize(distinctive) if distinctive else None

                    # 找出含"独特部分"的行（这些行 100% 是真命中），它们的 y_center 用来圈定"同一视觉行"
                    distinctive_y_centers = []
                    if distinctive_norm and distinctive_norm != orig_norm:
                        # 实体本身有独特部分（即名字含通用后缀）
                        for spec in line_specs:
                            lt_norm = normalize(spec[0])
                            if distinctive_norm in lt_norm:
                                yc = (spec[3] + spec[4]) / 2
                                distinctive_y_centers.append((yc, spec[4] - spec[3]))

                    for li, spec in enumerate(line_specs):
                        lt, lx0, lx1, ly0, ly1, gstart = spec
                        lt_norm = normalize(lt)
                        line_yc = (ly0 + ly1) / 2
                        line_h = ly1 - ly0

                        match = longest_common_substring(orig_norm, lt_norm, min_len=min_match_len)
                        if not match:
                            continue
                        ent_start, local_s, length = match
                        local_e = local_s + length
                        ent_end = ent_start + length

                        # 反误报关卡：
                        # 如果实体有独特部分（如公司有品牌名），且当前行不含独特部分，
                        # 则该行必须与某个含独特部分的行在同一视觉行（y 相近）才允许 redact
                        if distinctive_y_centers:
                            line_has_distinctive = distinctive_norm in lt_norm
                            if not line_has_distinctive:
                                tol = max(line_h * 0.6, 5)
                                same_row = any(
                                    abs(line_yc - dyc) <= tol
                                    for dyc, _ in distinctive_y_centers
                                )
                                if not same_row:
                                    continue  # 跳过：通用部分在不相干的行里出现，是误报

                        # 重叠检查
                        if any(s < local_e and local_s < e for s, e in line_handled[li]):
                            continue
                        line_handled[li].append((local_s, local_e))

                        line_w = max(lx1 - lx0, 1)
                        # 用 CJK 双倍宽权重精确算位置（避免英中混排时位置偏移）
                        sub_x0 = pixel_offset_in_line(lt, local_s, lx0, line_w)
                        sub_x1 = pixel_offset_in_line(lt, local_e, lx0, line_w)

                        pad_x, pad_y = 2, 2
                        rect_x0 = sub_x0 - pad_x
                        rect_x1 = sub_x1 + pad_x
                        rect_y0 = ly0 - pad_y
                        rect_y1 = ly1 + pad_y

                        draw.rectangle(
                            [(rect_x0, rect_y0), (rect_x1, rect_y1)],
                            fill="white", outline=None,
                        )

                        # 决定该段写什么文字：
                        #   - whitebox_only：强制不写任何文字，纯白框
                        #   - partial 掩码（mask 长度 == 实体长度）：写本行对应的切片（如"张*"）
                        #     这种掩码本身就有信息（"张*"显示有姓但不显示名），保留
                        #   - 占位符模式（如 [PERSON_1] / <人物1> / 〔姓名1〕）：在白框中央写完整占位符
                        if whitebox_only:
                            text_to_draw = ''
                        else:
                            same_length = len(masked) == len(orig_norm)
                            if same_length:
                                # partial 模式：本行只画对应切片
                                text_to_draw = masked[ent_start:ent_end]
                            else:
                                # 占位符模式：每个出现都画完整占位符（哪怕跨行也各画一份）
                                text_to_draw = masked

                        if text_to_draw:
                            line_h = ly1 - ly0
                            avail_w = rect_x1 - rect_x0
                            # 自适应字号：从行高 0.75 起，逐步缩小直到文字宽度 ≤ 95% 可用宽度
                            font_size = max(10, int(line_h * 0.75))
                            font = None
                            text_w = text_h = 0
                            for fs in range(font_size, 7, -1):
                                try:
                                    f = ImageFont.truetype(font_path, fs) if font_path else ImageFont.load_default()
                                    bb = draw.textbbox((0, 0), text_to_draw, font=f)
                                    tw = bb[2] - bb[0]
                                    th = bb[3] - bb[1]
                                except Exception:
                                    f = ImageFont.load_default()
                                    tw = fs * len(text_to_draw) // 2
                                    th = fs
                                if tw <= avail_w * 0.95:
                                    font, text_w, text_h = f, tw, th
                                    break
                            if font is None:
                                font = ImageFont.load_default()
                                text_w, text_h = avail_w * 0.9, line_h * 0.5

                            tx = rect_x0 + max((avail_w - text_w) / 2, 0)
                            ty = rect_y0 + max((rect_y1 - rect_y0 - text_h) / 2, 0)
                            draw.text((tx, ty), text_to_draw, fill="black", font=font)

                        page_redacted += 1

                # 把绘制完成的图像作为新页面背景
                new_page = out_doc.new_page(width=page.rect.width, height=page.rect.height)
                img_bytes = _io_bytes_png(img)
                new_page.insert_image(new_page.rect, stream=img_bytes)

                total_redacted += page_redacted
                elapsed = _time.time() - t0
                print(f"  第 {page_num}/{page_count} 页：脱敏 {page_redacted} 处，耗时 {elapsed:.1f}s", flush=True)

            print(f"[视觉脱敏] 完成，共脱敏 {total_redacted} 处", flush=True)

            out_doc.save(output_path, deflate=True, garbage=3)
            out_doc.close()
            doc.close()
            return True
        except Exception:
            import traceback
            traceback.print_exc()
            return False

    def write_file(self, text: str, output_path: str, output_format: str = 'auto') -> List[Tuple[str, str]]:
        """
        写入文件

        Args:
            text: 文本内容
            output_path: 输出文件路径
            output_format: 输出格式 (auto, txt, md, pdf, docx)

        Returns:
            列表 [(文件类型, 文件路径), ...]
        """
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)

        if output_format == 'auto':
            output_format = self._guess_format(path)

        saved_files = []

        if output_format == 'pdf':
            pdf_path = path if path.suffix.lower() == '.pdf' else path.with_suffix('.pdf')
            if self._write_pdf(text, str(pdf_path)):
                saved_files.append(('output_pdf', str(pdf_path)))
            else:
                # PDF生成失败，回退到文本
                txt_path = path.with_suffix('.txt')
                self._write_plain_text(text, str(txt_path))
                saved_files.append(('output_txt', str(txt_path)))

        elif output_format == 'docx':
            docx_path = path if path.suffix.lower() == '.docx' else path.with_suffix('.docx')
            if self._write_docx(text, str(docx_path)):
                saved_files.append(('output_docx', str(docx_path)))
            else:
                txt_path = path.with_suffix('.txt')
                self._write_plain_text(text, str(txt_path))
                saved_files.append(('output_txt', str(txt_path)))

        elif output_format == 'md':
            md_path = path if path.suffix.lower() == '.md' else path.with_suffix('.md')
            self._write_plain_text(text, str(md_path))
            saved_files.append(('output_md', str(md_path)))

        else:  # txt
            txt_path = path if path.suffix.lower() == '.txt' else path.with_suffix('.txt')
            self._write_plain_text(text, str(txt_path))
            saved_files.append(('output_txt', str(txt_path)))

        return saved_files

    def _guess_format(self, path: Path) -> str:
        """根据文件扩展名猜测格式"""
        suffix = path.suffix.lower()
        if suffix == '.pdf' and (HAS_REPORTLAB or HAS_PYMUPDF):
            return 'pdf'
        elif suffix in ['.docx', '.doc'] and HAS_DOCX:
            return 'docx'
        elif suffix == '.md':
            return 'md'
        else:
            return 'txt'

    def _write_plain_text(self, text: str, file_path: str) -> bool:
        """写入纯文本文件"""
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return True
        except Exception:
            return False

    def _write_pdf(self, text: str, file_path: str) -> bool:
        """写入PDF文件"""
        if HAS_REPORTLAB:
            try:
                return self._write_pdf_reportlab(text, file_path)
            except Exception:
                pass

        if HAS_PYMUPDF:
            try:
                return self._write_pdf_fitz(text, file_path)
            except Exception:
                pass

        return False

    def _write_pdf_reportlab(self, text: str, file_path: str) -> bool:
        """使用 reportlab 创建 PDF"""
        import platform
        import os

        doc = SimpleDocTemplate(str(file_path), pagesize=A4,
                                leftMargin=50, rightMargin=50,
                                topMargin=50, bottomMargin=50)
        story = []

        # 回退模板：仿宋 小四 1.5 倍行距
        # 小四 = 12pt，1.5x = 18pt leading
        styles = getSampleStyleSheet()
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=12,
            leading=18,  # 12pt * 1.5
            alignment=TA_LEFT,
            spaceBefore=0,
            spaceAfter=0,
        )

        # 尝试加载中文字体（支持 macOS / Windows / Linux）
        font_loaded = False
        font_paths = []
        if platform.system() == 'Darwin':
            # macOS 优先用仿宋类（Kaiti/STFangsong），否则回退 PingFang
            font_paths = [
                '/System/Library/Fonts/Supplemental/Songti.ttc',
                '/System/Library/Fonts/STHeiti Light.ttc',
                '/System/Library/Fonts/PingFang.ttc',
                '/System/Library/Fonts/Hiragino Sans GB.ttc',
            ]
        elif platform.system() == 'Windows':
            windir = os.environ.get('WINDIR', 'C:\\Windows')
            font_paths = [
                os.path.join(windir, 'Fonts', 'simfang.ttf'),   # 仿宋 ⭐优先
                os.path.join(windir, 'Fonts', 'simsun.ttc'),    # 宋体
                os.path.join(windir, 'Fonts', 'msyh.ttc'),      # 微软雅黑
                os.path.join(windir, 'Fonts', 'simhei.ttf'),    # 黑体
                os.path.join(windir, 'Fonts', 'msyhbd.ttc'),    # 微软雅黑粗体
            ]
        else:  # Linux
            font_paths = [
                '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
                '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
                '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
            ]

        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                    normal_style.fontName = 'ChineseFont'
                    font_loaded = True
                    break
                except Exception:
                    continue

        if not font_loaded:
            normal_style.fontName = 'Helvetica'

        lines = text.split('\n')
        for line in lines:
            if line.strip():
                p = Paragraph(line, normal_style)
                story.append(p)
            else:
                story.append(Spacer(1, 6))

        doc.build(story)
        return True

    def _write_pdf_fitz(self, text: str, file_path: str) -> bool:
        """使用 PyMuPDF 创建 PDF（降级方案，支持中文）"""
        import os

        doc = fitz.open()
        margin = 50
        font_size = 11
        line_height = 16
        page_height = 842
        page_width = 595
        max_y = page_height - margin

        # 尝试加载中文字体
        font_path = None
        font_name = "helv"

        # 跨平台中文字体搜索
        import platform as _plat
        candidate_fonts = [
            # macOS
            '/System/Library/Fonts/PingFang.ttc',
            '/System/Library/Fonts/STHeiti Light.ttc',
            '/System/Library/Fonts/Hiragino Sans GB.ttc',
            '/System/Library/Fonts/STSong.ttf',
            # Linux
            '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
            '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
            '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
        ]
        # Windows 字体
        if _plat.system() == 'Windows':
            windir = os.environ.get('WINDIR', 'C:\\Windows')
            candidate_fonts = [
                os.path.join(windir, 'Fonts', 'msyh.ttc'),
                os.path.join(windir, 'Fonts', 'simfang.ttf'),
                os.path.join(windir, 'Fonts', 'simsun.ttc'),
                os.path.join(windir, 'Fonts', 'simhei.ttf'),
            ] + candidate_fonts
        for fp in candidate_fonts:
            if os.path.exists(fp):
                font_path = fp
                break

        page = doc.new_page()
        current_y = margin

        if font_path:
            # 使用外部中文字体
            page.insert_font(fontname="cjk", fontfile=font_path)
            font_name = "cjk"

        lines = text.split('\n')
        for line in lines:
            if current_y + line_height > max_y:
                page = doc.new_page()
                current_y = margin
                if font_path:
                    page.insert_font(fontname="cjk", fontfile=font_path)

            if line.strip():
                try:
                    page.insert_text(
                        (margin, current_y), line,
                        fontsize=font_size, fontname=font_name
                    )
                except Exception:
                    # 最终降级：ASCII only
                    safe_line = ''.join([c if ord(c) < 128 else '?' for c in line])
                    page.insert_text(
                        (margin, current_y), safe_line,
                        fontsize=font_size, fontname="helv"
                    )

            current_y += line_height

        doc.save(file_path)
        doc.close()
        return True

    def _write_docx(self, text: str, file_path: str) -> bool:
        """
        写入 Word 文档 — 回退模板（仿宋 / 小四 / 1.5 倍行距）

        格式规范：
        - 字体：仿宋（中英文统一）
        - 正文：小四（12pt），行间距 1.5 倍，段前段后 0
        - 标题：小三（15pt）加粗，段前段后 0.5 行
        - 页边距：上 3.7cm 下 3.5cm 左 2.8cm 右 2.6cm
        """
        if not HAS_DOCX:
            return False

        try:
            from docx.oxml import OxmlElement
            from docx.shared import Cm, Emu
            from docx.enum.text import WD_LINE_SPACING
            import re as _re

            FONT_NAME = '仿宋'
            BODY_SIZE = Pt(12)       # 小四 = 12pt
            TITLE_SIZE = Pt(15)      # 小三 = 15pt
            LINE_SPACING = 1.5       # 1.5 倍行距
            TITLE_SPACE = Pt(6)      # 0.5 行 ≈ 6pt（基于 12pt 正文）

            doc = Document()

            # 页边距
            for section in doc.sections:
                section.top_margin = Cm(3.7)
                section.bottom_margin = Cm(3.5)
                section.left_margin = Cm(2.8)
                section.right_margin = Cm(2.6)

            def _set_rfonts_on_element(element):
                """为 XML 元素设置仿宋字体"""
                rPr = element.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    element.append(rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                for attr in ('w:eastAsia', 'w:ascii', 'w:hAnsi', 'w:cs'):
                    rFonts.set(qn(attr), FONT_NAME)
                # 清除主题字体覆盖
                for theme_attr in ('w:eastAsiaTheme', 'w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme'):
                    try:
                        del rFonts.attrib[qn(theme_attr)]
                    except (KeyError, ValueError):
                        pass

            def _set_line_spacing(para, spacing, rule=WD_LINE_SPACING.MULTIPLE):
                """设置段落行间距（默认 1.5 倍多行间距）"""
                pf = para.paragraph_format
                pf.line_spacing_rule = rule
                pf.line_spacing = spacing

            def _is_title(line):
                """判断是否为标题行（仅匹配文书名称和章节标题）"""
                stripped = line.strip()
                if not stripped or len(stripped) > 25:
                    return False
                # 文书名称（必须是独立的短标题）
                doc_titles = [
                    r'^(?:民事|刑事|行政)?(?:起诉状|答辩状|上诉状|代理意见|判决书|裁定书|调解书|决定书|申请书|异议书)$',
                    r'^(?:关于.+的(?:函|通知|公告|意见|决定|报告|说明|声明|承诺书))$',
                    r'^(?:租赁|买卖|借款|委托|合作|劳动|服务|股权转让)?(?:合同|协议)(?:书)?$',
                ]
                for pat in doc_titles:
                    if _re.search(pat, stripped):
                        return True
                # 章节标题（"诉讼请求" "事实与理由" 等独立短行）
                section_titles = [
                    r'^(?:诉讼请求|事实与理由|证据清单|判决如下|本院认为|裁判结果|审理查明)$',
                    r'^第[一二三四五六七八九十\d]+[章节部分](?:\s|$)',
                ]
                for pat in section_titles:
                    if _re.search(pat, stripped):
                        return True
                return False

            # 设置 Normal 样式默认值
            normal_style = doc.styles['Normal']
            normal_style.font.name = FONT_NAME
            normal_style.font.size = BODY_SIZE
            _set_rfonts_on_element(normal_style.element)
            normal_pf = normal_style.paragraph_format
            normal_pf.space_before = Pt(0)
            normal_pf.space_after = Pt(0)
            normal_pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            normal_pf.line_spacing = LINE_SPACING

            # 设置 docDefaults
            styles_element = doc.styles.element
            rPrDefault = styles_element.find(qn('w:docDefaults'))
            if rPrDefault is None:
                rPrDefault = OxmlElement('w:docDefaults')
                styles_element.insert(0, rPrDefault)
            rPrDef = rPrDefault.find(qn('w:rPrDefault'))
            if rPrDef is None:
                rPrDef = OxmlElement('w:rPrDefault')
                rPrDefault.append(rPrDef)
            rPr = rPrDef.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                rPrDef.append(rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            for attr in ('w:eastAsia', 'w:ascii', 'w:hAnsi', 'w:cs'):
                rFonts.set(qn(attr), FONT_NAME)
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz')
                rPr.append(sz)
            sz.set(qn('w:val'), '24')  # 12pt = 24 half-points
            szCs = rPr.find(qn('w:szCs'))
            if szCs is None:
                szCs = OxmlElement('w:szCs')
                rPr.append(szCs)
            szCs.set(qn('w:val'), '24')

            # 清洗：去掉 XML 不允许的控制字符（OCR 结果常见）
            # 只保留 \t \n \r 和 printable，丢弃其它 C0 控制字符
            import re as _re
            text = _re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

            # 处理文本内容
            lines = text.split('\n')
            for line in lines:
                stripped = line.strip()
                # 跳过页面分隔标记
                if stripped.startswith('=' * 10):
                    continue
                if stripped.startswith('第 ') and stripped.endswith(' 页'):
                    continue

                if stripped:
                    para = doc.add_paragraph()
                    is_title = _is_title(stripped)

                    run = para.add_run(stripped)
                    run.font.name = FONT_NAME
                    _set_rfonts_on_element(run._element)

                    if is_title:
                        run.font.size = TITLE_SIZE
                        run.bold = True
                        para.paragraph_format.space_before = TITLE_SPACE
                        para.paragraph_format.space_after = TITLE_SPACE
                    else:
                        run.font.size = BODY_SIZE
                    _set_line_spacing(para, LINE_SPACING)
                else:
                    # 空行：保持格式一致
                    para = doc.add_paragraph()
                    _set_rfonts_on_element(para._element)
                    _set_line_spacing(para, LINE_SPACING)

            doc.save(file_path)
            return True
        except Exception:
            import traceback
            traceback.print_exc()
            return False

    def write_mapping(self, mapping: Dict, file_path: str):
        """写入映射表"""
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(mapping, f, ensure_ascii=False, indent=2)

    def get_supported_input_formats(self) -> List[str]:
        """获取支持的输入格式"""
        return self.supported_formats['input'].copy()

    def get_supported_output_formats(self) -> List[str]:
        """获取支持的输出格式"""
        return self.supported_formats['output'].copy()
